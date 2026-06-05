from __future__ import annotations

from pathlib import Path

from docx import Document
from docx.enum.section import WD_SECTION
from docx.enum.table import WD_CELL_VERTICAL_ALIGNMENT, WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.shared import Cm, Inches, Pt, RGBColor
from docx.oxml import OxmlElement
from docx.oxml.ns import qn


ROOT = Path(__file__).resolve().parents[1]
OUT = ROOT / "DevDoc_De_Gouden_Draak_OTAP.docx"
SCREENSHOTS = ROOT / "devdoc_artifacts" / "screenshots"


def set_cell_shading(cell, fill: str) -> None:
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:fill"), fill)
    tc_pr.append(shd)


def set_cell_text(cell, text: str, bold: bool = False) -> None:
    cell.text = ""
    p = cell.paragraphs[0]
    p.paragraph_format.space_after = Pt(0)
    run = p.add_run(text)
    run.bold = bold
    run.font.name = "Calibri"
    run.font.size = Pt(9)


def add_table(doc: Document, headers: list[str], rows: list[list[str]], widths: list[float] | None = None):
    table = doc.add_table(rows=1, cols=len(headers))
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.style = "Table Grid"
    table.autofit = False
    hdr = table.rows[0].cells
    for i, header in enumerate(headers):
        set_cell_text(hdr[i], header, bold=True)
        set_cell_shading(hdr[i], "E8EEF5")
        hdr[i].vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
    for row in rows:
        cells = table.add_row().cells
        for i, value in enumerate(row):
            set_cell_text(cells[i], value)
            cells[i].vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
    if widths:
        for row in table.rows:
            for idx, width in enumerate(widths):
                row.cells[idx].width = Cm(width)
    doc.add_paragraph()
    return table


def add_bullet(doc: Document, text: str) -> None:
    p = doc.add_paragraph(style="List Bullet")
    p.add_run(text)


def add_number(doc: Document, text: str) -> None:
    p = doc.add_paragraph(style="List Number")
    p.add_run(text)


def add_code(doc: Document, text: str) -> None:
    p = doc.add_paragraph()
    p.paragraph_format.left_indent = Cm(0.5)
    p.paragraph_format.space_after = Pt(4)
    run = p.add_run(text)
    run.font.name = "Courier New"
    run.font.size = Pt(8)
    run.font.color.rgb = RGBColor(40, 40, 40)


def add_source(doc: Document, caption: str, path: str) -> None:
    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(2)
    run = p.add_run(f"Bewijs: {caption} - {path}")
    run.italic = True
    run.font.size = Pt(9)
    run.font.color.rgb = RGBColor(80, 80, 80)


def style_document(doc: Document) -> None:
    section = doc.sections[0]
    section.top_margin = Inches(0.85)
    section.bottom_margin = Inches(0.85)
    section.left_margin = Inches(0.85)
    section.right_margin = Inches(0.85)

    styles = doc.styles
    normal = styles["Normal"]
    normal.font.name = "Calibri"
    normal.font.size = Pt(10.5)
    normal.paragraph_format.space_after = Pt(6)
    normal.paragraph_format.line_spacing = 1.08

    for style_name, size, color, before, after in [
        ("Heading 1", 16, "1F4D78", 14, 6),
        ("Heading 2", 13, "2E74B5", 10, 4),
        ("Heading 3", 11.5, "1F4D78", 8, 3),
    ]:
        st = styles[style_name]
        st.font.name = "Calibri"
        st.font.size = Pt(size)
        st.font.bold = True
        st.font.color.rgb = RGBColor.from_string(color)
        st.paragraph_format.space_before = Pt(before)
        st.paragraph_format.space_after = Pt(after)


def add_title_page(doc: Document) -> None:
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_after = Pt(0)
    run = p.add_run("DevDoc")
    run.font.name = "Calibri"
    run.font.size = Pt(28)
    run.font.bold = True
    run.font.color.rgb = RGBColor.from_string("1F4D78")

    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run("De Gouden Draak - OTAP, migratie en modernisering")
    run.font.size = Pt(16)
    run.font.color.rgb = RGBColor.from_string("2E74B5")

    doc.add_paragraph()
    add_table(
        doc,
        ["Onderdeel", "Invulling"],
        [
            ["Project", "Update van de verouderde website van restaurant De Gouden Draak"],
            ["Techniek nieuw", "Laravel, VueJS, Apache, MySQL/MariaDB"],
            ["Techniek oud", "Statische HTML, PHP-kassa, MySQL/MariaDB SQL-dump"],
            ["Datum analyse", "2 juni 2026"],
            ["Doel document", "Onderbouwde technische basis voor bouw, OTAP, deployment en datamigratie"],
        ],
        [4, 12],
    )

    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.add_run("Status: eerste complete Word-versie voor review en export naar PDF").italic = True
    doc.add_page_break()


def add_scope(doc: Document) -> None:
    doc.add_heading("1. Samenvatting en uitgangspunten", level=1)
    doc.add_paragraph(
        "De eigenaar van De Gouden Draak wil de bestaande website moderniseren zonder dat de huidige site tijdens de ontwikkeling offline gaat. "
        "De nieuwe oplossing moet overdraagbaar zijn aan zijn zonen, die Laravel en VueJS leren. Daarom wordt gekozen voor een Laravel-backend met VueJS-frontend, "
        "draaiend op een omgeving die de shared hosting van de huidige provider zo dicht mogelijk nabootst: Apache, PHP en MySQL/MariaDB."
    )
    doc.add_paragraph(
        "De belangrijkste ontwerpkeuze is dat de nieuwe site niet als losse pagina-update wordt behandeld, maar als schaalbare applicatie met gescheiden omgevingen, "
        "database-migraties, versiebeheer, deploymentafspraken en een gecontroleerd datamigratieproces. Daardoor kan de bestaande menukaart, kassa-data en toekomstige besteldata beheersbaar worden overgezet."
    )
    doc.add_heading("Concrete randvoorwaarden uit de opdracht", level=2)
    for item in [
        "De bestaande productieomgeving blijft online tot acceptatie is afgerond.",
        "De nieuwe applicatie gebruikt Laravel voor backend, routing, validatie, authenticatie en datalaag.",
        "De frontend gebruikt VueJS voor interactieve onderdelen zoals kassa, bestelzuil en menuconfigurator.",
        "Apache en MySQL/MariaDB worden als doelplatform aangehouden, omdat de huidige shared hosting dit aanbiedt.",
        "Er is geen directe SSH-toegang op productie; beheer vindt plaats via hostingpanel, phpMyAdmin en file upload/deploy-export.",
        "Er mag geen bestaande functionaliteit of data verloren gaan.",
    ]:
        add_bullet(doc, item)


def add_current_functionality(doc: Document) -> None:
    doc.add_heading("2. Inventarisatie huidige functionaliteit", level=1)
    doc.add_paragraph(
        "De oude website bestaat uit statische HTML-pagina's en een aparte PHP-kassa. De navigatie biedt menukaart, nieuws en contact. "
        "De kassa gebruikt een database met menu-items en registreert verkochte items in de tabel sales."
    )
    add_table(
        doc,
        ["Gebied", "Huidige functionaliteit", "Bewijs"],
        [
            ["Publieke site", "Homepage met branding, studentaanbieding en navigatie naar menukaart, nieuws en contact.", "webfs_old/index.html"],
            ["Menukaart", "HTML-pagina met links naar menukaart-PDF/JPG-bestanden.", "webfs_old/paginas/MENUKAART.html"],
            ["Nieuws", "Nieuwsbericht over beperkte opening tijdens de coronacrisis.", "webfs_old/paginas/news.html"],
            ["Contact", "Contactpagina met route-informatie en afbeelding.", "webfs_old/paginas/contact.html"],
            ["Kassa-login", "Medewerker logt in met medewerker-nummer en wachtwoord.", "webfs_old/kassa/loginRequest.php"],
            ["Kassa-bestelling", "Medewerker selecteert menu-items en rekent bestelling af.", "webfs_old/kassa/cashDesk.php en payOrder.php"],
            ["Verkoopoverzicht", "Medewerker bekijkt verkopen binnen een datumselectie.", "webfs_old/kassa/sales.php en salesOverview.php"],
        ],
        [3.2, 8.2, 4.5],
    )

    doc.add_heading("User stories huidige situatie", level=2)
    stories = [
        [
            "US-01",
            "Als bezoeker wil ik de homepage van De Gouden Draak bekijken, zodat ik direct zie welk restaurant dit is en welke aanbieding actief is.",
            "Homepage toont naam, restauranttype, studentaanbieding en navigatie naar menukaart, nieuws en contact.",
        ],
        [
            "US-02",
            "Als bezoeker wil ik de menukaart openen, zodat ik gerechten en prijzen kan bekijken voordat ik bestel.",
            "Menukaartpagina is bereikbaar via navigatie en toont of linkt naar de restaurantmenukaart.",
        ],
        [
            "US-03",
            "Als bezoeker wil ik nieuws en openinginformatie bekijken, zodat ik weet of het restaurant beschikbaar is.",
            "Nieuwspagina toont actuele mededelingen en is bereikbaar via navigatie.",
        ],
        [
            "US-04",
            "Als bezoeker wil ik contact- en route-informatie bekijken, zodat ik het restaurant kan vinden.",
            "Contactpagina toont routeinformatie en is bereikbaar via navigatie.",
        ],
        [
            "US-05",
            "Als medewerker wil ik inloggen op de kassa, zodat alleen personeel bestellingen kan verwerken.",
            "Login accepteert medewerker-nummer en wachtwoord en geeft een foutmelding bij lege of onbekende combinatie.",
        ],
        [
            "US-06",
            "Als medewerker wil ik menu-items aan een bestelling toevoegen, zodat ik een afhaalbestelling kan samenstellen.",
            "Kassa toont menu-items per soort gerecht, laat aantallen wijzigen en berekent het totaalbedrag.",
        ],
        [
            "US-07",
            "Als medewerker wil ik een bestelling afrekenen, zodat de verkoop wordt opgeslagen.",
            "Afrekenen schrijft per item een regel naar sales met itemId, amount en saleDate.",
        ],
        [
            "US-08",
            "Als eigenaar wil ik verkopen per periode bekijken, zodat ik inzicht heb in omzet en populaire gerechten.",
            "Verkoopoverzicht filtert sales op begindatum en einddatum.",
        ],
    ]
    add_table(doc, ["ID", "User story", "Acceptatiecriteria"], stories, [1.7, 7.7, 7.0])


def add_new_functionality_from_excel(doc: Document) -> None:
    doc.add_heading("3. Nieuwe functionaliteit uit bijlage", level=1)
    doc.add_paragraph(
        "De nieuwe functionaliteiten zijn gebaseerd op de omgezette bijlage WEBFS - Nieuwe Functionaliteit v2.xlsx. "
        "De bijlage bevat user stories met acceptatiecriteria en storypoints. De afgesproken scope voor deze opdracht is: US-4, US-6, US-7, US-9, US-10, UC-14, UC-15, UC-19 en UC-20. "
        "Deze selectie haalt de knockout-eis, omdat er minimaal twee features uit elke categorie worden gekozen: Restaurant/tablet, Kassa voor afhaal, Website en Admin."
    )
    add_table(
        doc,
        ["Categorie", "Aantal stories", "Minimaal voor KO", "Advies voor eerste release"],
        [
            ["Restaurant/tablet", "8", "2", "US-4 rekening als PDF, US-6 herhalingsbestelling en US-7 cocktail-inspiratie."],
            ["Kassa voor afhaal", "4", "2", "US-9 zoek/filter en US-10 opmerkingen per gerecht."],
            ["Website", "4 plus responsive requirement", "2", "UC-14 gegenereerde menu-PDF en UC-15 favoriete gerechten."],
            ["Admin", "4", "2", "UC-19 dagelijkse verkooprapportage en UC-20 menu CRUD."],
        ],
        [4.0, 3.0, 3.0, 6.0],
    )

    doc.add_page_break()
    doc.add_heading("Geselecteerde scope", level=2)
    selected_rows = [
        ["US-4", "Restaurant/tablet", "Rekening opslaan als PDF", "PDF-bewijs met logo, meerdere pagina's en productregels."],
        ["US-6", "Restaurant/tablet", "Herhalingsbestelling vanuit historie", "Orderhistorie en opnieuw plaatsen van geheel of delen van vorige bestelling."],
        ["US-7", "Restaurant/tablet", "Cocktail-inspiratie via externe API", "Zoeken, filteren en afbeeldingen tonen uit TheCocktailDB of vergelijkbare API."],
        ["US-9", "Kassa voor afhaal", "Gerechten zoeken en filteren", "Zoeken op naam/nummer en filteren op categorie."],
        ["US-10", "Kassa voor afhaal", "Opmerkingen per gerecht", "Vrije opmerkingen en suggesties uit eerder gebruikte opmerkingen."],
        ["UC-14", "Website", "Menu-PDF genereren vanuit database", "Actuele menu-PDF inclusief aanbiedingen."],
        ["UC-15", "Website", "Favoriete gerechten via cookie", "Favorieten bovenaan tonen en sorteeropties ondersteunen."],
        ["UC-19", "Admin", "Dagelijkse verkooprapportage", "Laravel scheduled task maakt Excel-rapportage downloadbaar."],
        ["UC-20", "Admin", "Gerechten CRUD", "Menu beheren met behoud van bestaande nummering en a/b/c-varianten."],
    ]
    add_table(doc, ["ID", "Categorie", "Feature", "Waarom in scope"], selected_rows, [1.6, 3.5, 4.3, 7.0])

    grouped_stories = {
        "Restaurant (tablet view)": [
            [
                "US-1",
                "Als klant wil ik een bestelling kunnen maken aan mijn tafel door middel van een tablet zodat ik rustig de tijd kan nemen om te kiezen.",
                "Tafelindicatie aanwezig. Kan maar 1x per 10 minuten bestellen. Maximaal 5 rondes.",
                "5",
            ],
            [
                "US-2",
                "Als ober wil ik bezoekers registreren bij een tafel zodat ik de rekening daarop kan aanpassen.",
                "Maximaal 8 personen per tafel. Leeftijd van de personen registreren. Extra DELUXE menu ja/nee.",
                "7",
            ],
            [
                "US-3",
                "Als ober wil ik de rekening kunnen opdelen waarbij je per deel producten kan koppelen.",
                "Maximaal 8 delen. Verschillende aantallen van producten kunnen per deel worden aangegeven, bijvoorbeeld 3 bier opdelen in drie delen.",
                "4",
            ],
            [
                "US-4",
                "Als ober wil ik de rekening kunnen opslaan als PDF.",
                "Logo van het restaurant. Leesbare indeling. Test met minimaal 35 producten en meerdere pagina's. Afbeelding, naam, stukprijs, aantal en totaalprijs staan op de rekening. Formaat 8,5 cm breed en 10 cm lang per pagina.",
                "4",
            ],
            [
                "US-5",
                "Als klant wil ik hulp kunnen inschakelen zodat een ober naar mijn tafel komt; als ober wil ik alle hulpvragen zien en afmelden.",
                "In de backoffice verschijnt een melding. Ober kan hulpvraag afmelden na hulpverlening.",
                "4",
            ],
            [
                "US-6",
                "Als klant wil ik een herhalingsbestelling doen bij mijn tafel zodat ik makkelijk opnieuw kan bestellen vanuit historie.",
                "Klant kan de gehele vorige bestelling opnieuw doen of delen van de vorige bestelling herhalen.",
                "6",
            ],
            [
                "US-7",
                "Als klant wil ik tijdens het wachten cocktail-inspiratie zien zodat ik een drankje kan bestellen.",
                "Gebruik een cocktail-API zoals TheCocktailDB. Toon afbeeldingen. Bied zoekfunctie en filterfunctie.",
                "6",
            ],
            [
                "US-8",
                "De rekening bevat een QR-code die naar een responsive reviewformulier wijst.",
                "Reviewformulier bevat meerdere vragen en is interactief/responsive.",
                "10",
            ],
        ],
        "Kassa voor afhaal (desktop view)": [
            [
                "US-9",
                "Als kassamedewerker kan ik eenvoudig gerechten opzoeken zodat klanten sneller geholpen worden.",
                "Invoermogelijkheid met zoekfunctie. Zoeken op gerechtnaam en gerechtnummer. Filteren op gerechtcategorie.",
                "4",
            ],
            [
                "US-10",
                "Als kassamedewerker kan ik opmerkingen toevoegen aan een gerecht op een bestelling zodat klantspecifieke wensen worden meegenomen.",
                "Ieder gerecht op een bestelling kan opmerkingen krijgen, bijvoorbeeld geen ui. Er is een lijst met meest gebruikte opmerkingen uit voorgaande bestellingen.",
                "3",
            ],
            [
                "US-11",
                "Als kassamedewerker kan ik per gerecht aangeven of de klant witte rijst, nasi/bami goreng, mihoen goreng of Chinese bami wil.",
                "Rijstkeuze kan direct bij het gerecht worden gekozen. Meerdere nasi/bami-bestellingen gaan via een nieuwe orderregel.",
                "4",
            ],
            [
                "US-12",
                "Als kassamedewerker kan ik aanbiedingen invoeren voor de aankomende week.",
                "Aanbiedingen zijn zichtbaar op een aparte websitepagina en op het gegenereerde menu. Een aanbieding is een prijsaanpassing op een of meerdere producten voor een bepaalde periode.",
                "10",
            ],
        ],
        "Website": [
            [
                "UC-13",
                "Als klant lees ik de website in een geschikte taal zodat ik de site kan begrijpen.",
                "Localisatie in frontend/backend. Site toont automatisch Nederlands of Engels. Handmatig wisselen tussen NL en EN. Geen twee fysiek aparte sites. Vertaalteksten zijn aanpasbaar via JSON of database. Controle op veelvoorkomende accessibility-problemen.",
                "8",
            ],
            [
                "UC-14",
                "Als klant wil ik een up-to-date PDF van het menu kunnen downloaden.",
                "PDF wordt gegenereerd vanuit de database. Eventuele aanbiedingen worden op een aparte pagina in de PDF afgedrukt.",
                "4",
            ],
            [
                "UC-15",
                "Als klant wil ik favoriete gerechten kunnen aanvinken zodat mijn voorkeur lokaal bewaard blijft.",
                "Favorieten worden als cookie opgeslagen. Menu kan favorieten bovenaan tonen op nummer. Alleen favorieten kunnen alfabetisch worden gesorteerd terwijl de rest van het menu gelijk blijft.",
                "4",
            ],
            [
                "UC-16",
                "Als klant wil ik via de website een afhaalbestelling doorgeven.",
                "Bedanktpagina toont printbare QR-code met bestelnummer, gerechtnummers en gerechtnamen.",
                "4",
            ],
            [
                "REQ-W01",
                "De website is responsive op drie verschillende schermgroottes.",
                "Layout werkt op mobiel, tablet en desktop zonder horizontaal scrollen of overlappende elementen.",
                "7",
            ],
        ],
        "Admin": [
            [
                "UC-17",
                "Als restaurant-eigenaar kan ik een planning maken voor maximaal een week zodat ik medewerkers kan koppelen aan meerdere tafels.",
                "Personen zijn beschikbaar om te koppelen zonder uitgebreide user/login CRUD. Planning koppelt medewerkers aan tafels.",
                "11",
            ],
            [
                "UC-18",
                "Er is een CMS aanwezig om op de klantwebsite het menu in te stellen en extra pagina's te maken.",
                "CMS ondersteunt vaste layout en WYSIWYG-editor voor extra pagina's en menubeheer.",
                "11",
            ],
            [
                "UC-19",
                "Als admin krijg ik dagelijks automatisch een samenvatting van de verkoop van die dag.",
                "Samenvatting wordt automatisch gegenereerd met Laravel task scheduling. Output is Excel-format. Samenvattingen staan op een webpagina en zijn downloadbaar.",
                "8",
            ],
            [
                "UC-20",
                "Als admin kan ik gerechten toevoegen, aanpassen en verwijderen zodat het menu up-to-date blijft.",
                "Nummering blijft gelijk. Nieuwe varianten gebruiken a/b/c bij bestaande nummers. Nieuwe gerechten of categorieen mogen doornummeren.",
                "6",
            ],
        ],
    }

    for category, rows in grouped_stories.items():
        doc.add_heading(category, level=2)
        add_table(doc, ["ID", "User story", "Acceptatiecriteria", "SP"], rows, [1.6, 6.4, 7.4, 1.0])

    doc.add_heading("Impact op architectuur en datamodel", level=2)
    add_table(
        doc,
        ["Onderdeel", "Stories", "Technische consequentie"],
        [
            ["Rekening als PDF", "US-4", "Orders moeten PDF-generatie, restaurantlogo, meerdere pagina's en orderregel-totalen ondersteunen."],
            ["Herhalingsbestelling", "US-6", "Orders en order_lines moeten historisch beschikbaar blijven en als template voor nieuwe order gebruikt kunnen worden."],
            ["Cocktail-inspiratie", "US-7", "Frontend krijgt API-client, zoek/filter-state en caching/rate-limit-afspraken."],
            ["Kassa zoeken/filteren", "US-9", "Menu-items krijgen indexen op naam, nummer en categorie; frontend krijgt snelle zoekcomponent."],
            ["Opmerkingen per gerecht", "US-10", "Order_lines krijgen notes; note_suggestions worden afgeleid uit eerdere orderregels."],
            ["Menu-PDF", "UC-14", "PDF-generatie leest menu_categories, menu_items en promotions vanuit de database."],
            ["Favorieten", "UC-15", "Client-side cookie/localStorage voor favorieten; backend hoeft geen klantaccount te maken."],
            ["Dagrapportage", "UC-19", "Laravel scheduler genereert Excel-bestanden en bewaart report_files voor download."],
            ["Menu CRUD", "UC-20", "Admin CRUD rond menu_categories en menu_items met validatie op nummering en suffixen."],
        ],
        [4.0, 3.6, 8.6],
    )


def add_database(doc: Document) -> None:
    doc.add_heading("4. Database-inventarisatie oude situatie", level=1)
    doc.add_paragraph(
        "De oude database heet gouden_draak. De dump is gemaakt met MySQL dump 10.13 en vermeldt serverversie 5.5.5-10.1.37-MariaDB. "
        "De database gebruikt standaard latin1. Dit is technisch risicovol voor moderne meertalige data en moet bij migratie naar utf8mb4 worden omgezet."
    )
    add_table(
        doc,
        ["Tabel", "Doel", "Belangrijke kolommen", "Constatering"],
        [
            ["gebruiker", "Kassa-login voor medewerkers", "id, wachtwoord, isAdmin", "Wachtwoord staat plaintext in de dump: test. Geen naam, rolmodel of hashing."],
            ["menu", "Menukaart en kassa-items", "menunummer, menu_toevoeging, naam, price, soortgerecht, beschrijving", "Hoofddata van de menukaart. price is float en categorie staat als tekst in elke rij."],
            ["menu_pdf", "Opslag van menukaart-PDF", "id, menu_pdf", "Lege tabel; BLOB-opslag is niet wenselijk voor beheer op shared hosting."],
            ["rijst_enzo", "Keuzes/toeslagen voor rijst/bami/etc.", "beschrijving, extra_prijs", "Lege tabel; bedoeld voor configuratie maar niet gevuld."],
            ["sales", "Verkoopregels uit kassa", "itemId, amount, saleDate", "Geen order-header. Eén bestelling wordt niet als geheel opgeslagen, alleen losse regels."],
            ["specialiteiten", "Specialiteiten met rijstkeuze", "id, name, informatie, prijs, rijst_keuze_etc", "Lege tabel met foreign key naar lege rijst_enzo-tabel."],
        ],
        [2.5, 4.3, 5.3, 5.0],
    )
    doc.add_heading("Datakwaliteit en technische fouten", level=2)
    for item in [
        "Karakterset latin1 moet naar utf8mb4 om Nederlandse, Chinese en speciale tekens betrouwbaar te bewaren.",
        "Prijzen in menu.price gebruiken float; geldbedragen moeten decimal(8,2) worden om afrondingsfouten te voorkomen.",
        "Gebruikerswachtwoorden zijn niet gehasht. In Laravel moet Hash::make gebruikt worden.",
        "SQL in loginRequest.php en salesOverview.php wordt opgebouwd met stringconcatenatie. Dit is gevoelig voor SQL-injectie.",
        "sales.itemId heeft geen expliciete foreign key naar menu.id in de dump.",
        "Een bestelling heeft geen ordernummer, betaalstatus, kanaal, tafel/zuil of klanttype. Daardoor is latere synchronisatie en rapportage beperkt.",
        "Menu-categorieen staan als vrije tekst in menu.soortgerecht, waardoor typefouten en inconsistente categorievolgorde mogelijk zijn.",
        "Sommige gerechtteksten bevatten HTML zoals <br> en entiteiten zoals &eacute;; deze moeten bij migratie opgeschoond worden.",
    ]:
        add_bullet(doc, item)

    doc.add_heading("Voorstel nieuw datamodel", level=2)
    add_table(
        doc,
        ["Nieuwe tabel", "Waarom", "Belangrijke velden"],
        [
            ["users", "Medewerkers en beheerders via Laravel-authenticatie.", "id, name, email, password, role, active"],
            ["menu_categories", "Schaalbare categorie-indeling voor menukaart en kassa.", "id, name, sort_order, active"],
            ["menu_items", "Vervangt menu met genormaliseerde velden.", "id, category_id, number, suffix, name, description, price, active"],
            ["promotions", "Aanbiedingen kunnen in menu-PDF en website worden meegenomen.", "id, title, starts_at, ends_at, discount_type, active"],
            ["promotion_items", "Koppelt aanbiedingen aan een of meerdere gerechten.", "id, promotion_id, menu_item_id, discount_amount"],
            ["orders", "Bestelling als geheel, nodig voor PDF-rekening, herhaalbestelling en rapportage.", "id, channel, table_code, status, subtotal, total, paid_at, created_at"],
            ["order_lines", "Bestelregels per gerecht.", "id, order_id, menu_item_id, quantity, unit_price, line_total"],
            ["order_line_notes", "Opmerkingen per gerecht en herbruikbare suggesties.", "id, order_line_id, note, normalized_note"],
            ["generated_files", "Opslag van gegenereerde PDF-menu's, rekeningen en Excel-rapportages.", "id, type, path, generated_at, generated_by"],
            ["favorite_menu_items", "Optionele serverkant voor statistiek; klantfavorieten blijven primair cookie/localStorage.", "id, menu_item_id, count"],
            ["api_cache", "Cache voor cocktail-API-resultaten om externe API minder te belasten.", "id, source, cache_key, payload, expires_at"],
            ["content_pages", "Beheerbare content voor nieuws/contact zonder codewijziging.", "id, slug, title, body, published_at"],
        ],
        [3.2, 6.0, 6.0],
    )


def add_otap(doc: Document) -> None:
    doc.add_heading("5. OTAP-inrichting", level=1)
    doc.add_paragraph(
        "De OTAP-straat bestaat uit Ontwikkel, Test, Acceptatie en Productie. Iedere omgeving krijgt een eigen database, eigen .env-configuratie en eigen applicatie-URL. "
        "Omdat productie shared hosting zonder SSH is, wordt productie niet als ontwikkelomgeving gebruikt. Alleen gevalideerde releases worden naar productie gebracht."
    )
    add_table(
        doc,
        ["Omgeving", "URL", "Doel", "Database", "Toegang"],
        [
            ["Ontwikkel", "http://goudendraak.test", "Lokale ontwikkeling door team", "gouden_draak_dev", "Lokale laptop via Git"],
            ["Test", "https://test.goudendraak.example", "Automatische en handmatige tests", "gouden_draak_test", "Team en docent"],
            ["Acceptatie", "https://acc.goudendraak.example", "Goedkeuring door eigenaar en zonen", "gouden_draak_acc", "Eigenaar, zonen, team"],
            ["Productie", "https://www.degoudendraak.example", "Live website en kassa", "gouden_draak_prod", "Hostingpanel en phpMyAdmin"],
        ],
        [2.4, 4.4, 4.0, 3.4, 3.4],
    )
    doc.add_heading("Gewenste serverconfiguratie", level=2)
    add_table(
        doc,
        ["Onderdeel", "Keuze", "Onderbouwing"],
        [
            ["Webserver", "Apache 2.4 met mod_rewrite", "Past bij shared hosting en Laravel public/.htaccess-routing."],
            ["PHP", "PHP 8.4 of 8.5; minimaal 8.3 voor Laravel 13", "PHP 8.4 en 8.5 hebben op 2 juni 2026 actieve support volgens php.net. Laravel 13 ondersteunt PHP 8.3-8.5."],
            ["Database", "MySQL 8.0/8.4 of MariaDB 10.11+", "Modernere versie dan oude MariaDB 10.1.37, met utf8mb4, foreign keys en betere indexering."],
            ["Node build", "Node.js 22 LTS voor buildomgeving", "Alleen nodig voor Vite/Vue build; niet noodzakelijk op shared productie als assets vooraf gebouwd worden."],
            ["Composer", "Composer 2.x", "Nodig voor Laravel dependencies tijdens build of lokale releasevoorbereiding."],
            ["Opslag", "Minimaal 2 GB voor applicatie, uploads, exports en backups", "Menukaarten, logs en exports groeien; opslag moet niet op limiet zitten."],
            ["PHP-extensies", "pdo_mysql, mbstring, openssl, tokenizer, xml, ctype, json, fileinfo, curl", "Nodig voor Laravel, database, validatie, uploads en HTTP-integraties."],
        ],
        [3.2, 4.5, 8.5],
    )
    doc.add_heading("Apache virtual host instructie", level=2)
    doc.add_paragraph("Gebruik voor elke niet-productieomgeving een eigen DocumentRoot naar de public-map van Laravel. Voor acceptatie:")
    add_code(doc, "<VirtualHost *:80>\n    ServerName acc.goudendraak.example\n    DocumentRoot /var/www/gouden-draak/current/public\n\n    <Directory /var/www/gouden-draak/current/public>\n        AllowOverride All\n        Require all granted\n    </Directory>\n\n    ErrorLog ${APACHE_LOG_DIR}/gouden-draak-acc-error.log\n    CustomLog ${APACHE_LOG_DIR}/gouden-draak-acc-access.log combined\n</VirtualHost>")
    doc.add_paragraph(
        "Op productie zonder SSH wordt dezelfde public-map-structuur aangehouden via het hostingpanel: de publieke documentroot moet naar de Laravel public-map verwijzen. "
        "Als de provider dit niet toestaat, wordt de inhoud van public in de webroot geplaatst en verwijzen index.php en assetpaden naar de applicatiemap buiten de webroot."
    )
    doc.add_heading("Per-omgeving .env-afspraken", level=2)
    add_code(doc, "APP_ENV=acceptance\nAPP_DEBUG=false\nAPP_URL=https://acc.goudendraak.example\nDB_CONNECTION=mysql\nDB_HOST=localhost\nDB_DATABASE=gouden_draak_acc\nDB_USERNAME=gouden_draak_acc_user\nDB_PASSWORD=<uniek wachtwoord per omgeving>")


def add_version_control(doc: Document) -> None:
    doc.add_heading("6. Versiebeheer", level=1)
    doc.add_paragraph(
        "Versiebeheer wordt toegepast met Git. De repository bevat Laravel-code, Vue-code, database-migraties, seeders, tests en documentatie. "
        ".env-bestanden, database-exports, uploads en productiebackups worden niet gecommit."
    )
    add_table(
        doc,
        ["Branch", "Doel", "Regel"],
        [
            ["main", "Productiewaardige code", "Alleen via pull request en na acceptatie."],
            ["develop", "Integratie van afgeronde features", "Automatisch deploybaar naar Test."],
            ["feature/<naam>", "Nieuwe functionaliteit", "Korte levensduur, gekoppeld aan user story."],
            ["release/<versie>", "Acceptatie en stabilisatie", "Alleen bugfixes, geen nieuwe features."],
            ["hotfix/<naam>", "Spoedfix productie", "Vanaf main, daarna terugmergen naar develop."],
        ],
        [3.3, 5.2, 7.2],
    )
    doc.add_heading("Commit- en tagafspraken", level=2)
    for item in [
        "Commitberichten beschrijven wat is aangepast, bijvoorbeeld: feat(kassa): voeg orderregels toe.",
        "Elke databasewijziging gaat via Laravel migration; directe handmatige tabelwijzigingen zijn alleen toegestaan als noodprocedure en worden daarna als migration vastgelegd.",
        "Elke release naar Acceptatie krijgt een tag zoals v1.0.0-rc.1.",
        "Elke productie-release krijgt een tag zoals v1.0.0 met release-notes en rollback-instructie.",
    ]:
        add_bullet(doc, item)


def add_deployment(doc: Document) -> None:
    doc.add_heading("7. Deploymentplan en opleverafspraken", level=1)
    doc.add_paragraph(
        "Deployments worden gepland om dataverlies en downtime te voorkomen. De oude productie blijft actief tot de eigenaar de acceptatieomgeving heeft goedgekeurd. "
        "Nieuwe functionaliteit wordt eerst op Test gecontroleerd en daarna op Acceptatie beoordeeld."
    )
    add_table(
        doc,
        ["Moment", "Afspraak", "Reden"],
        [
            ["Feature freeze", "5 werkdagen voor geplande oplevering naar Acceptatie.", "Vanaf dit moment geen nieuwe wensen meer in dezelfde release."],
            ["Content freeze", "24 uur voor productiemigratie.", "Voorkomt dat nieuws/menuwijzigingen na de laatste export verloren gaan."],
            ["Database freeze", "Tijdens de definitieve migratie, maximaal 30 minuten.", "Voorkomt verschillen tussen oude en nieuwe database."],
            ["Deploymoment", "Buiten piekuren, bijvoorbeeld maandag 09:00-10:00.", "Restaurantverkeer is dan lager dan rond avondeten/weekend."],
            ["Backup", "Voor elke productie-release: bestanden en database exporteren.", "Rollback moet uitvoerbaar zijn."],
            ["Synchronisatie", "Productiedata nooit overschrijven met testdata.", "Beschermt echte bestellingen en menu-aanpassingen."],
        ],
        [3.1, 7.1, 5.5],
    )
    doc.add_heading("Stappenplan productie-oplevering", level=2)
    steps = [
        "Zet release branch vast en maak tag v1.0.0-rc.1.",
        "Deploy naar Acceptatie en voer regressietest uit op homepage, menukaart, login, kassa, menuconfigurator en bestelproces.",
        "Laat eigenaar en zonen acceptatie uitvoeren met vooraf afgesproken scenario's.",
        "Start content freeze en exporteer oude database via phpMyAdmin.",
        "Maak backup van oude bestanden en oude database.",
        "Voer migratiescript uit op de nieuwe productiedatabase.",
        "Controleer aantallen menu-items, gebruikers en verkoopregels na migratie.",
        "Upload releasepakket naar productie en controleer APP_ENV=production en APP_DEBUG=false.",
        "Voer smoke test uit: homepage, menukaart, login, bestelling aanmaken, verkoopoverzicht.",
        "Hef content freeze op en noteer release in changelog.",
    ]
    for step in steps:
        add_number(doc, step)
    doc.add_heading("Rollback", level=2)
    doc.add_paragraph(
        "Als de smoke test faalt en de fout niet binnen 15 minuten veilig opgelost kan worden, wordt teruggerold naar de oude bestanden en oude databasebackup. "
        "Nieuwe orders die eventueel al in de nieuwe database zijn geplaatst, worden eerst als SQL-export veiliggesteld voordat de oude database wordt teruggezet."
    )


def add_data_management(doc: Document) -> None:
    doc.add_heading("8. Data management en synchronisatie", level=1)
    doc.add_paragraph(
        "Data stroomt gecontroleerd van productie naar lagere omgevingen voor testdoeleinden, en code stroomt van ontwikkel naar productie. "
        "Er wordt nooit testdata teruggezet naar productie."
    )
    add_table(
        doc,
        ["Richting", "Toegestaan?", "Instructie"],
        [
            ["Productie -> Acceptatie/Test", "Ja, na anonimisering", "Exporteer via phpMyAdmin, anonimiseer klantgegevens en importeer in lagere omgeving."],
            ["Acceptatie/Test -> Productie", "Nee", "Alleen Laravel migrations en gecontroleerde seeders mogen naar productie."],
            ["Oude DB -> Nieuwe DB", "Ja, via migratiescript", "Map menu naar menu_categories/menu_items en sales naar orders/order_lines."],
            ["Uploads/menu-PDF -> Nieuwe storage", "Ja", "Plaats bestanden in storage/app/public of hosting uploadmap en registreer metadata in database."],
        ],
        [4, 3, 9],
    )
    doc.add_heading("Technische migratie-instructie oude data", level=2)
    for item in [
        "Exporteer gouden_draak via phpMyAdmin als SQL met structuur en data.",
        "Importeer de dump lokaal in gouden_draak_legacy.",
        "Maak een Laravel migration voor de nieuwe tabellen en voer php artisan migrate uit op gouden_draak_dev.",
        "Maak een Artisan command app:import-legacy-gouden-draak dat de oude tabellen leest en de nieuwe tabellen vult.",
        "Converteer latin1-tekst naar utf8mb4 en verwijder HTML uit naamvelden waar mogelijk.",
        "Maak menu_categories op basis van DISTINCT soortgerecht, met sort_order volgens de oude menukaartvolgorde.",
        "Maak menu_items aan met decimal-prijzen en behoud oud id in legacy_menu_id voor controle.",
        "Groepeer sales-regels met dezelfde saleDate tot orders wanneer de timestamps gelijk zijn; zet losse regels in order_lines.",
        "Controleer aantallen: 1 gebruiker, 175 menu-items minus ontbrekende id 166, 35 sales-regels en lege configuratietabellen.",
    ]:
        add_bullet(doc, item)
    doc.add_heading("Controlequeries na migratie", level=2)
    add_code(doc, "SELECT COUNT(*) FROM menu;\nSELECT COUNT(*) FROM menu_items;\nSELECT COUNT(*) FROM sales;\nSELECT COUNT(*) FROM order_lines;\nSELECT legacy_menu_id, name FROM menu_items WHERE legacy_menu_id IS NULL;")


def add_frontend(doc: Document) -> None:
    doc.add_heading("9. Frontendadvies", level=1)
    doc.add_paragraph(
        "Het advies is om Vue 3 met TypeScript te gebruiken binnen de Laravel/Vite-stack. Dit sluit aan op de opdracht, is overdraagbaar aan studenten die Laravel en VueJS leren, "
        "en maakt interactieve onderdelen schaalbaar zonder de backend te mengen met complexe frontendstatus."
    )
    add_table(
        doc,
        ["Keuze", "Advies", "Onderbouwing"],
        [
            ["Framework", "Vue 3 + TypeScript", "Vue is verplicht vanuit de opdracht en TypeScript verlaagt foutkans bij kassa en bestelconfiguratie."],
            ["Laravel-integratie", "Inertia.js of Laravel API + Vue SPA", "Inertia is geschikt voor snelle Laravel/Vue-overdracht; API+SPA is geschikt als tablets later losser moeten draaien."],
            ["Buildtool", "Vite", "Standaard in moderne Laravel en snel voor development/builds."],
            ["Styling", "Tailwind CSS of Bootstrap 5", "Tailwind geeft maatwerk zonder oude tabel-layout; Bootstrap is eenvoudiger voor studenten. Advies: Tailwind als team dit beheerst."],
            ["State management", "Pinia alleen waar nodig", "Kassa en menuconfigurator hebben gedeelde winkelmand-status; simpele pagina's niet."],
            ["Validatie", "Laravel Form Requests + Vue client-validatie", "Server blijft leidend; frontend geeft snelle feedback."],
            ["Testing", "Pest/PHPUnit voor backend, Vitest voor Vue componenten", "Dek kritieke flows zoals bestellen, prijzen en migratiecontrole af."],
        ],
        [3.4, 5, 8],
    )
    doc.add_heading("Nieuwe functionaliteit", level=2)
    add_table(
        doc,
        ["Functionaliteit", "Aanpak", "Datamodel-impact"],
        [
            ["Rekening-PDF", "Laravel genereert een compacte PDF-rekening vanuit orderdata; Vue toont preview/downloadstatus.", "orders, order_lines, order_pdfs"],
            ["Herhalingsbestelling", "Vue toont vorige bestellingen en laat de klant een volledige order of losse regels opnieuw toevoegen.", "orders, order_lines, reorder_source_id"],
            ["Cocktail-inspiratie", "Vue-component haalt cocktails op via externe API met zoek- en filterfunctie.", "Geen kern-DB nodig; optioneel api_cache"],
            ["Kassa zoeken/filteren", "Desktopgerichte Vue-kassa met zoekveld, categoriefilter en snelle selectie.", "menu_categories, menu_items indexen"],
            ["Opmerkingen per gerecht", "Orderregel krijgt vrij tekstveld en suggesties uit eerder gebruikte opmerkingen.", "order_lines.notes, note_suggestions"],
            ["Menu-PDF", "Admin kan actuele PDF genereren vanuit database inclusief aanbiedingen.", "menu_items, promotions, generated_files"],
            ["Favorieten", "Website bewaart favorieten in cookie/localStorage en sorteert menuweergave.", "Geen account nodig; menu_items blijven brondata"],
            ["Dagrapportage", "Laravel scheduled task maakt dagelijks Excel-overzicht voor admin-download.", "report_files, orders, order_lines"],
            ["Menu CRUD", "Admin beheert gerechten, categorieen, prijzen en a/b/c-nummering.", "menu_categories, menu_items"],
        ],
        [3.4, 8.0, 4.7],
    )


def add_evidence(doc: Document) -> None:
    doc.add_heading("10. Bewijsstukken", level=1)
    doc.add_paragraph(
        "Deze bijlage toont bewijs van de oude opdrachtbestanden. De afbeeldingen zijn Quick Look-thumbnails van de lokale HTML-pagina's. "
        "Daarnaast zijn scriptbestanden en SQL-bestanden als bewijs vermeld, omdat de KO screenshots of script files toestaat."
    )
    images = [
        ("Homepage oude website", SCREENSHOTS / "index.html.png"),
        ("Menukaartpagina oude website", SCREENSHOTS / "MENUKAART.html.png"),
        ("Nieuwspagina oude website", SCREENSHOTS / "news.html.png"),
        ("Contactpagina oude website", SCREENSHOTS / "contact.html.png"),
    ]
    for caption, path in images:
        if path.exists():
            p = doc.add_paragraph()
            p.add_run(caption).bold = True
            doc.add_picture(str(path), width=Inches(5.8))
            add_source(doc, caption, str(path.relative_to(ROOT)))
    doc.add_heading("Belangrijke scriptbestanden", level=2)
    add_table(
        doc,
        ["Bestand", "Bewijswaarde"],
        [
            ["WEBFS - Nieuwe Functionaliteit v2.xlsx", "Omgezette bijlage met nieuwe user stories, acceptatiecriteria en storypoints."],
            ["webfs_old/gouden_draak_create_script.sql", "Database-schema, menu-data, sales-data en oude MariaDB-versie."],
            ["webfs_old/kassa/config/dbconfig.php", "Hardcoded lokale databaseconfiguratie."],
            ["webfs_old/kassa/loginRequest.php", "Login-flow en SQL-query op gebruiker."],
            ["webfs_old/kassa/cashDesk.php", "Menu-items worden uit database gelezen en als kassalijst getoond."],
            ["webfs_old/kassa/payOrder.php", "Bestelling wordt als sales-regels opgeslagen."],
            ["webfs_old/kassa/salesOverview.php", "Verkoopoverzicht per datumselectie."],
        ],
        [6, 10],
    )
    doc.add_heading("Codefragmenten", level=2)
    add_code(doc, "define(\"DB_SERVERNAME\", \"localhost\");\ndefine(\"DB_DBNAME\", \"gouden_draak\");\ndefine(\"DB_USERNAME\", \"root\");\ndefine(\"DB_PASSWORD\", \"\");")
    add_source(doc, "Hardcoded databaseconfiguratie", "webfs_old/kassa/config/dbconfig.php")
    add_code(doc, "SELECT * FROM gebruiker WHERE id='\".$_POST['employeeNr'].\"' AND wachtwoord='\".$_POST['password'].\"'")
    add_source(doc, "Login-query zonder hashing/prepared statements", "webfs_old/kassa/loginRequest.php")
    add_code(doc, "INSERT INTO sales (itemId, amount, saleDate) VALUES ($order->id, $order->amount, NOW())")
    add_source(doc, "Opslaan van kassa-verkoopregels", "webfs_old/kassa/payOrder.php")


def add_rubric(doc: Document) -> None:
    doc.add_heading("11. Rubric-checklist", level=1)
    add_table(
        doc,
        ["Rubricpunt", "Waar afgedekt", "Status"],
        [
            ["Algemeen: onderbouwing, naamgeving, netheid", "Alle hoofdstukken bevatten keuze-onderbouwing; tabelnamen zijn technisch consistent.", "Aanwezig"],
            ["Functionaliteit: user stories", "Hoofdstuk 2 met oude user stories en hoofdstuk 3 met alle nieuwe stories uit de Excel-bijlage.", "Aanwezig"],
            ["OTAP", "Hoofdstuk 5 met URL's, databases, Apache config, .env en versies.", "Aanwezig"],
            ["Versiebeheer", "Hoofdstuk 6 met branches, tags, migrations en release-afspraken.", "Aanwezig"],
            ["Deployment", "Hoofdstuk 7 met freeze, backup, deploymoment, smoke test en rollback.", "Aanwezig"],
            ["Database", "Hoofdstuk 4 met oude tabellen, datakwaliteit en nieuw datamodel.", "Aanwezig"],
            ["Data management", "Hoofdstuk 8 met synchronisatierichting, migratie-instructie en controlequeries.", "Aanwezig"],
            ["Frontend", "Hoofdstuk 9 met Vue/Laravel advies en packages.", "Aanwezig"],
            ["Bewijs", "Hoofdstuk 10 met screenshots/thumbnails, Excel-bijlage en scriptbestanden.", "Aanwezig"],
            ["PDF-inlevering", "Exporteer dit Word-document na review naar PDF.", "Nog uitvoeren bij inlevering"],
        ],
        [5.6, 8.0, 2.4],
    )
    doc.add_heading("Bronnen voor versieadvies", level=2)
    for item in [
        "PHP Supported Versions, geraadpleegd op 2 juni 2026: https://www.php.net/supported-versions.php",
        "Laravel Release Notes en Support Policy, geraadpleegd op 2 juni 2026: https://laravel.com/docs/master/releases",
        "npm package vue, geraadpleegd op 2 juni 2026: https://www.npmjs.com/package/vue",
    ]:
        add_bullet(doc, item)


def main() -> None:
    doc = Document()
    style_document(doc)
    add_title_page(doc)
    add_scope(doc)
    add_current_functionality(doc)
    add_new_functionality_from_excel(doc)
    add_database(doc)
    add_otap(doc)
    add_version_control(doc)
    add_deployment(doc)
    add_data_management(doc)
    add_frontend(doc)
    add_evidence(doc)
    add_rubric(doc)
    doc.save(OUT)
    print(OUT)


if __name__ == "__main__":
    main()
